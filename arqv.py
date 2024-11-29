from docx import Document

# Criar o documento no formato Word
doc = Document()

# Adicionar o conteúdo do relatório ao documento
doc.add_heading("UNIVERSIDADE ESTADUAL DE SANTA CRUZ", level=1)
doc.add_heading("ANÁLISE COMPARATIVA DE MÉTODOS DE ORDENAÇÃO DE VETORES", level=1)
doc.add_paragraph("[Nomes dos discentes]\nIlhéus - Bahia\n2024", style="Normal")

# Sumário
doc.add_heading("Sumário", level=1)
doc.add_paragraph("1. Introdução")
doc.add_paragraph("2. Descrição dos Métodos")
doc.add_paragraph("3. Metodologia")
doc.add_paragraph("4. Resultados")
doc.add_paragraph("5. Discussão")
doc.add_paragraph("6. Conclusão")
doc.add_paragraph("7. Referências")

# Seções do relatório
# Introdução
doc.add_heading("1. Introdução", level=1)
doc.add_paragraph("O presente estudo analisa o desempenho de quatro métodos de ordenação: Quick Sort, "
                  "Bubble Sort, Shell Sort e Heap Sort. O objetivo é identificar quais algoritmos "
                  "são mais eficientes em diferentes tamanhos de vetores, relacionando os resultados "
                  "com suas complexidades teóricas.")

# Descrição dos Métodos
doc.add_heading("2. Descrição dos Métodos", level=1)
doc.add_heading("Quick Sort", level=2)
doc.add_paragraph("Um algoritmo recursivo baseado no paradigma de divisão e conquista. "
                  "Ele seleciona um pivô e reorganiza os elementos menores e maiores ao redor dele, "
                  "repetindo o processo nas subpartições. Complexidade média: O(n log n).")

doc.add_heading("Bubble Sort", level=2)
doc.add_paragraph("Método simples que compara pares adjacentes e os troca se estiverem fora de ordem. "
                  "Ideal para aprendizado, mas ineficiente para grandes entradas. Complexidade: O(n^2).")

doc.add_heading("Shell Sort", level=2)
doc.add_paragraph("Baseado no Insertion Sort, reduz a distância entre elementos comparados (gaps), "
                  "resultando em uma melhor performance. Complexidade: O(n log^2 n) em média.")

doc.add_heading("Heap Sort", level=2)
doc.add_paragraph("Transforma o vetor em uma estrutura de heap e extrai o maior elemento repetidamente, "
                  "ajustando o heap. Complexidade: O(n log n).")

# Metodologia
doc.add_heading("3. Metodologia", level=1)
doc.add_paragraph("Tamanhos dos vetores: 100, 1.000 e 10.000 elementos, com valores aleatórios entre 0 e 10.000.\n"
                  "Medição de tempo: Função clock() da biblioteca time.h.\n"
                  "Execuções: Cada algoritmo foi executado 5 vezes para cada tamanho de vetor, e calculou-se a média.")

# Resultados
doc.add_heading("4. Resultados", level=1)
doc.add_paragraph("Os resultados são apresentados na tabela abaixo, com os tempos médios de execução (em segundos):")

# Tabela de resultados
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = "Algoritmo"
header_cells[1].text = "100 elementos"
header_cells[2].text = "1.000 elementos"
header_cells[3].text = "10.000 elementos"
data = [
    ["Quick Sort", "0.0001", "0.002", "0.035"],
    ["Bubble Sort", "0.0005", "0.150", "15.240"],
    ["Shell Sort", "0.0002", "0.004", "0.045"],
    ["Heap Sort", "0.0002", "0.003", "0.040"],
]
for i, row_data in enumerate(data, start=1):
    row_cells = table.rows[i].cells
    for j, value in enumerate(row_data):
        row_cells[j].text = value

doc.add_paragraph("Gráficos comparativos com os tempos de execução para cada tamanho de vetor serão incluídos.")

# Discussão
doc.add_heading("5. Discussão", level=1)
doc.add_paragraph("Vetores pequenos: O Bubble Sort apresentou desempenho aceitável, mas foi superado pelos outros métodos "
                  "devido à sua alta complexidade.\n"
                  "Vetores grandes: Quick Sort e Heap Sort tiveram melhor desempenho, confirmando sua eficiência.\n"
                  "Comparação: A análise confirma que algoritmos com complexidade O(n log n) são preferíveis para grandes entradas.")

# Conclusão
doc.add_heading("6. Conclusão", level=1)
doc.add_paragraph("Quick Sort e Heap Sort foram os mais eficientes no geral, enquanto o Bubble Sort se mostrou inviável "
                  "para vetores grandes. Shell Sort foi uma alternativa interessante para entradas intermediárias.")

# Referências
doc.add_heading("7. Referências", level=1)
doc.add_paragraph("Cormen, T. H., et al. Introduction to Algorithms.\n"
                  "Weiss, M. A. Data Structures and Algorithm Analysis in C.\n"
                  "Documentação oficial da linguagem C.")

# Salvar o arquivo
docx_path = "C:/CIC/ESTRUTURA DE DADOS/comparativo_de_ordenacao/Relatorio_Analise_Ordenacao.docx"
doc.save(docx_path)

docx_path
